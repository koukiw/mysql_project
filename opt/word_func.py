from docx import Document
import os
import re
import sys
import csv
import json
import glob
import shutil
import warnings

#これからやりたいこと
# 表ごとにjson形式でデータを格納
# 添付画像がある場合はそれも格納、ある場合のみjsonのkeyを増やす

# wordからテキストを抜き出す
def word2text(collection_name,dt_now):
    process_name = ""
    try:
        #
        # PDF→テキスト情報抽出処理
        #
        process_name = "wordからのテキスト情報抽出"
        # files = glob.glob(os.path.join("./word", "*.docx"))
        # files = glob.glob(os.path.join("./file_dir/**", "*.docx"),recursive=True)
        if collection_name =="":
            files = glob.glob(os.path.join("./file_dir","*.docx"),recursive=True)
        else:
            files = glob.glob(os.path.join("./file_dir", collection_name,"**","*.docx"),recursive=True)
        print("wordfiles",files)
        cnt = 0
        Input_path = "./word"
        results = []
        for file in files:
            cnt += 1
            print("wordテキスト抽出処理中…（{}/{}）".format(cnt, len(files)))
            # print(file)

            # テキストの抜き出し
            # 本文中のテキスト抽出
            text = ""
            document = Document(file)
            for i, p in enumerate(document.paragraphs):
                text += p.text

            # 表中のテキスト抽出
            for i, t in enumerate(document.tables):
                for j, r in enumerate(t.rows):
                    for k, c in enumerate(r.cells):
                        text += c.text
            # テキストの加工
            # text = text.replace("\n","").replace("\r","").replace("\t","").strip()

            filename  = os.path.basename(file)

            result = {"プロジェクト名":collection_name,"パス":file[11:],"ファイル名":filename,"text": text,"ファイル拡張子":"word","upload_data":dt_now}
            results.append(result)

        return results
            
    except Exception as e:
        print("error発生")
        return -1



# メイン処理
if __name__ == "__main__":

    results = word2text()
    print(results)

